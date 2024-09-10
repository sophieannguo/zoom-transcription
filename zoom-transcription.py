"""Turns a file with auto-transcribed text from Zoom meetings into a more readable Word document."""
import re
from docx import Document


def read_docx(file_path):
    """Reads text from a Word document."""
    doc = Document(file_path)
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    return '\n'.join(full_text)


def clean_and_combine_text(text):
    """Cleans and organizes the transcription text by speaker."""
    # Regular expression to remove lines where a number is followed by a timestamp
    removed_text = re.sub(r'\d+\s+\d{2}:\d{2}:\d{2}\.\d{3} --> \d{2}:\d{2}:\d{2}\.\d{3}', '', text)

    # Initialize a list to store the speech in order
    ordered_speech = []

    # Split the cleaned text by lines
    lines = removed_text.splitlines()

    # Variable to track the last speaker
    last_speaker = None

    for line in lines:
        # Match lines that contain a speaker, allowing for commas, hyphens, slashes, and apostrophes in names
        match = re.match(r'([\w\s,\-\/\']+): (.+)', line)
        if match:
            speaker, speech = match.groups()
            speech = speech.strip()

            # Check if the speaker is the same as the last one
            if speaker == last_speaker:
                # Combine speech with the last entry
                ordered_speech[-1] = f"{speaker}: {ordered_speech[-1].split(': ', 1)[1]} {speech}"
            else:
                # Add a new entry for a new speaker
                ordered_speech.append(f"{speaker}: {speech}")
                last_speaker = speaker

    # Create the final cleaned text
    cleaned_text = "\n".join(ordered_speech)

    return cleaned_text.strip()


def save_to_docx(clean_text, new_path):
    """Saves the given text to a Word document with line breaks after each speaker."""
    doc = Document()

    # Split the text into lines based on newlines
    paragraphs = clean_text.split('\n')

    for paragraph in paragraphs:
        if paragraph.strip():
            doc.add_paragraph(paragraph)
            # Add an extra line break after each speaker's speech
            doc.add_paragraph()  # This creates an empty paragraph to add a line break

    doc.save(new_path)


# Example usage
input_path = r"C:\Users\guophie\Downloads\Filename.vtt.docx"  # Adjust as needed
output_path = r"C:\Users\guophie\Downloads\Cleaned_Filename.docx"  # Path for the new document

input_text = read_docx(input_path)
final_text = clean_and_combine_text(input_text)
save_to_docx(final_text, output_path)

print(f"Cleaned text has been saved to {output_path}")